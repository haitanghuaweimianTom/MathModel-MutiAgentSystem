"""
文档加载器
==========

支持多种文档格式的统一加载接口。
"""

import os
from pathlib import Path
from typing import List, Optional, Dict, Any, Callable
from dataclasses import dataclass, field
from enum import Enum


class DocumentType(str, Enum):
    """文档类型"""
    EXCEL = "excel"
    WORD = "word"
    PDF = "pdf"
    MARKDOWN = "markdown"
    TEXT = "text"
    CSV = "csv"
    UNKNOWN = "unknown"


@dataclass
class DocumentContent:
    """文档内容"""
    filepath: str
    filename: str
    doc_type: DocumentType
    text: str = ""  # 文本内容
    tables: List[Any] = field(default_factory=list)  # 表格数据
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


@dataclass
class LoadResult:
    """加载结果"""
    success: bool
    document: Optional[DocumentContent] = None
    error: Optional[str] = None


class DocumentLoader:
    """文档加载器"""

    # 文件扩展名到文档类型的映射
    EXTENSION_MAP = {
        ".xlsx": DocumentType.EXCEL,
        ".xls": DocumentType.EXCEL,
        ".docx": DocumentType.WORD,
        ".pdf": DocumentType.PDF,
        ".md": DocumentType.MARKDOWN,
        ".txt": DocumentType.TEXT,
        ".csv": DocumentType.CSV,
    }

    def __init__(self):
        self._handlers: Dict[DocumentType, Callable[[str], DocumentContent]] = {
            DocumentType.EXCEL: self._load_excel,
            DocumentType.WORD: self._load_word,
            DocumentType.PDF: self._load_pdf,
            DocumentType.MARKDOWN: self._load_text,
            DocumentType.TEXT: self._load_text,
            DocumentType.CSV: self._load_csv,
        }

    def detect_type(self, filepath: str) -> DocumentType:
        """根据文件扩展名检测文档类型"""
        ext = Path(filepath).suffix.lower()
        return self.EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)

    def load(self, filepath: str) -> LoadResult:
        """
        加载单个文档

        Args:
            filepath: 文件路径

        Returns:
            LoadResult: 加载结果
        """
        path = Path(filepath)
        if not path.exists():
            return LoadResult(
                success=False,
                error=f"文件不存在: {filepath}"
            )

        doc_type = self.detect_type(filepath)
        if doc_type == DocumentType.UNKNOWN:
            return LoadResult(
                success=False,
                error=f"不支持的文件格式: {path.suffix}"
            )

        handler = self._handlers.get(doc_type)
        if not handler:
            return LoadResult(
                success=False,
                error=f"未找到对应的加载器: {doc_type.value}"
            )

        try:
            doc = handler(filepath)
            if doc.error:
                return LoadResult(success=False, document=doc, error=doc.error)
            return LoadResult(success=True, document=doc)
        except Exception as e:
            return LoadResult(
                success=False,
                error=f"加载失败: {str(e)}"
            )

    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        include_types: Optional[List[DocumentType]] = None,
    ) -> List[DocumentContent]:
        """
        批量加载目录中的文档

        Args:
            directory: 目录路径
            recursive: 是否递归子目录
            include_types: 仅加载指定类型的文档

        Returns:
            List[DocumentContent]: 加载成功的文档列表
        """
        dir_path = Path(directory)
        if not dir_path.exists():
            print(f"[DocumentLoader] 目录不存在: {directory}")
            return []

        if recursive:
            files = list(dir_path.rglob("*"))
        else:
            files = list(dir_path.iterdir())

        results = []
        for file_path in files:
            if not file_path.is_file():
                continue

            doc_type = self.detect_type(str(file_path))
            if doc_type == DocumentType.UNKNOWN:
                continue

            if include_types and doc_type not in include_types:
                continue

            result = self.load(str(file_path))
            if result.success and result.document:
                results.append(result.document)
            else:
                print(f"[DocumentLoader] 加载失败 {file_path}: {result.error}")

        print(f"[DocumentLoader] 从 {directory} 加载了 {len(results)} 个文档")
        return results

    def load_to_knowledge_base(
        self,
        directory: str,
        knowledge_base,
        recursive: bool = True,
    ) -> int:
        """
        将目录中的文档加载到知识库

        Returns:
            int: 成功加载的文档数量
        """
        from src.knowledge import KnowledgeBase

        documents = self.load_directory(directory, recursive=recursive)
        count = 0
        for doc in documents:
            try:
                knowledge_base.add_document(
                    title=doc.filename,
                    content=doc.text,
                    source=doc.filepath,
                    metadata={
                        "type": doc.doc_type.value,
                        **doc.metadata,
                    }
                )
                count += 1
            except Exception as e:
                print(f"[DocumentLoader] 添加到知识库失败 {doc.filename}: {e}")

        return count

    # =================================================================
    # 各格式加载实现
    # =================================================================

    def _load_excel(self, filepath: str) -> DocumentContent:
        """加载 Excel 文件"""
        try:
            import pandas as pd
            df = pd.read_excel(filepath)

            # 生成文本描述
            lines = [
                f"文件名: {Path(filepath).name}",
                f"行数: {len(df)}",
                f"列数: {len(df.columns)}",
                f"列名: {list(df.columns)}",
                "",
                "数据预览 (前10行):",
                df.head(10).to_string(),
                "",
                "数据统计:",
                df.describe().to_string(),
            ]

            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.EXCEL,
                text="\n".join(lines),
                tables=[df.to_dict(orient="records")],
                metadata={"rows": len(df), "columns": list(df.columns)},
            )
        except ImportError:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.EXCEL,
                error="pandas/openpyxl 未安装，无法读取 Excel 文件",
            )
        except Exception as e:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.EXCEL,
                error=str(e),
            )

    def _load_word(self, filepath: str) -> DocumentContent:
        """加载 Word 文件"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)

            # 提取表格
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append(rows)

            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.WORD,
                text=text,
                tables=tables,
                metadata={
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                },
            )
        except ImportError:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.WORD,
                error="python-docx 未安装，请运行: pip install python-docx",
            )
        except Exception as e:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.WORD,
                error=str(e),
            )

    def _load_pdf(self, filepath: str) -> DocumentContent:
        """加载 PDF 文件"""
        # 尝试多种 PDF 库
        errors = []

        # 尝试 PyPDF2
        try:
            import PyPDF2
            text_parts = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            text = "\n\n".join(text_parts)
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.PDF,
                text=text,
                metadata={"pages": len(reader.pages)},
            )
        except ImportError:
            errors.append("PyPDF2")
        except Exception as e:
            errors.append(f"PyPDF2: {e}")

        # 尝试 pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            text = "\n\n".join(text_parts)
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.PDF,
                text=text,
                metadata={"pages": len(reader.pages)},
            )
        except ImportError:
            errors.append("pypdf")
        except Exception as e:
            errors.append(f"pypdf: {e}")

        # 尝试 pdfplumber
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            text = "\n\n".join(text_parts)
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.PDF,
                text=text,
                metadata={"pages": len(text_parts)},
            )
        except ImportError:
            errors.append("pdfplumber")
        except Exception as e:
            errors.append(f"pdfplumber: {e}")

        return DocumentContent(
            filepath=filepath,
            filename=Path(filepath).name,
            doc_type=DocumentType.PDF,
            error=f"PDF 读取失败。尝试的库: {', '.join(errors)}。"
                  f"请安装其中之一: pip install PyPDF2 pypdf pdfplumber",
        )

    def _load_text(self, filepath: str) -> DocumentContent:
        """加载文本文件"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="gbk") as f:
                text = f.read()

        doc_type = DocumentType.MARKDOWN if filepath.endswith(".md") else DocumentType.TEXT
        return DocumentContent(
            filepath=filepath,
            filename=Path(filepath).name,
            doc_type=doc_type,
            text=text,
            metadata={"chars": len(text), "lines": text.count("\n") + 1},
        )

    def _load_csv(self, filepath: str) -> DocumentContent:
        """加载 CSV 文件"""
        try:
            import pandas as pd
            df = pd.read_csv(filepath)

            lines = [
                f"文件名: {Path(filepath).name}",
                f"行数: {len(df)}",
                f"列数: {len(df.columns)}",
                f"列名: {list(df.columns)}",
                "",
                "数据预览 (前10行):",
                df.head(10).to_string(),
            ]

            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.CSV,
                text="\n".join(lines),
                tables=[df.to_dict(orient="records")],
                metadata={"rows": len(df), "columns": list(df.columns)},
            )
        except ImportError:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.CSV,
                error="pandas 未安装，无法读取 CSV 文件",
            )
        except Exception as e:
            return DocumentContent(
                filepath=filepath,
                filename=Path(filepath).name,
                doc_type=DocumentType.CSV,
                error=str(e),
            )
