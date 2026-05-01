"""
统一工作流引擎 v2.0
===================

融合 LLM-MM-Agent + cherry-studio + Claude CLI 的统一架构：
- Coordinator: DAG调度与黑板内存（LLM-MM-Agent）
- CritiqueEngine: Actor-Critic-Improvement 质量保障（LLM-MM-Agent）
- KnowledgeBase: RAG 知识增强（cherry-studio）
- DocumentLoader: 多格式文档处理（cherry-studio）
- CodeExecutor: Claude CLI 代码生成与执行
- PaperGenerator: 大纲驱动的论文章节生成（LLM-MM-Agent）
- PaperTemplate: 通用论文模板系统

核心改进：
1. 每章节完整生成、内容充实、可直接交付
2. 代码任务明确使用 Claude CLI
3. 数据分析利用 KnowledgeBase + DocumentLoader
4. 支持多种论文模板（数学建模/课程作业/金融分析）
"""

import os
import sys
import json
import time
import re
import shutil
import subprocess
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from enum import Enum

import numpy as np
import pandas as pd

# =============================================================================
# 多 LLM 提供商支持
# =============================================================================

try:
    from src.llm import get_provider_manager, ProviderType, LLMProviderFactory, ProviderConfig
    _LLM_PROVIDER_AVAILABLE = True
except ImportError:
    _LLM_PROVIDER_AVAILABLE = False


_llm_provider_instance = None


def _get_llm_provider():
    """获取 LLM Provider 实例（惰性初始化）"""
    global _llm_provider_instance
    if _llm_provider_instance is not None:
        return _llm_provider_instance

    if not _LLM_PROVIDER_AVAILABLE:
        return None

    try:
        if os.getenv("OPENAI_API_KEY"):
            manager = get_provider_manager()
            manager.register(ProviderType.OPENAI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 OpenAI Provider (model={manager.get().config.model})")
        elif os.getenv("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_AUTH_TOKEN"):
            manager = get_provider_manager()
            config = ProviderConfig.from_env(ProviderType.ANTHROPIC)
            if not config.api_key and os.getenv("ANTHROPIC_AUTH_TOKEN"):
                config.api_key = os.getenv("ANTHROPIC_AUTH_TOKEN")
            if os.getenv("ANTHROPIC_BASE_URL"):
                config.api_host = os.getenv("ANTHROPIC_BASE_URL")
            manager.register(ProviderType.ANTHROPIC, config)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Anthropic Provider (model={manager.get().config.model})")
        elif os.getenv("GEMINI_API_KEY"):
            manager = get_provider_manager()
            manager.register(ProviderType.GEMINI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Gemini Provider (model={manager.get().config.model})")
        elif os.getenv("OLLAMA_MODEL") or os.getenv("OLLAMA_HOST"):
            manager = get_provider_manager()
            manager.register(ProviderType.OLLAMA)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Ollama Provider (model={manager.get().config.model})")
        else:
            manager = get_provider_manager()
            manager.register(ProviderType.CLAUDE_CLI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Claude CLI Provider (model={manager.get().config.model})")
    except Exception as e:
        print(f"[LLM] Provider 初始化失败: {e}，将回退到 Claude CLI")
        _llm_provider_instance = None

    return _llm_provider_instance


def _find_claude_code() -> Optional[str]:
    """自动搜索 Claude Code CLI 路径"""
    found = shutil.which("claude-code") or shutil.which("claude")
    return found


def _call_llm(
    prompt: str,
    system_prompt: Optional[str] = None,
    model: str = "sonnet",
    timeout: int = 600,
    max_retries: int = 3,
    retry_wait: int = 5
) -> str:
    """
    调用 LLM 生成内容
    优先使用配置的 API Provider，回退到 Claude Code CLI
    """
    provider_manager = _get_llm_provider()

    if provider_manager is not None:
        for attempt in range(max_retries):
            try:
                if attempt > 0:
                    print(f"    [LLM调用 {attempt + 1}/{max_retries}] 等待{retry_wait}秒后重试...")
                    time.sleep(retry_wait)

                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 开始请求...")
                response = provider_manager.generate(
                    prompt=prompt,
                    system_prompt=system_prompt,
                    timeout=timeout
                )
                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 成功!")
                return response
            except Exception as e:
                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 失败: {str(e)[:200]}")
                if attempt < max_retries - 1:
                    continue
                print("    [LLM] API Provider 失败，回退到 Claude CLI...")
                break

    # 回退到 Claude Code CLI
    claude_path = _find_claude_code()
    if not claude_path:
        raise RuntimeError(
            "Claude Code CLI 未找到，请确保已安装 Claude Code 并添加到 PATH，"
            "或配置 OPENAI_API_KEY / ANTHROPIC_API_KEY / GEMINI_API_KEY 环境变量"
        )

    full_prompt = prompt
    if system_prompt:
        full_prompt = f"{system_prompt}\n\n{prompt}"

    cmd = [
        claude_path,
        "-p",
        "--model", model,
        "--output-format", "json",
        full_prompt
    ]

    for attempt in range(max_retries):
        try:
            if attempt > 0:
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 等待{retry_wait}秒后重试...")
                time.sleep(retry_wait)

            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 开始请求...")
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            stdout, stderr = proc.communicate(timeout=timeout)

            stdout_text = stdout.decode("utf-8", errors="replace").strip()

            if proc.returncode != 0:
                error_msg = stderr.decode("utf-8", errors="replace")
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 失败: {error_msg[:200]}")
                if attempt < max_retries - 1:
                    continue
                raise RuntimeError(f"LLM 调用失败: {error_msg[:500]}")

            try:
                data = json.loads(stdout_text.strip())
            except json.JSONDecodeError:
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
                return stdout_text.strip()

            result_text = data.get("result", "")
            if isinstance(result_text, str):
                result_text = result_text.strip()
                if result_text.startswith("```"):
                    lines = result_text.splitlines()
                    if lines:
                        if lines[0].startswith("```"):
                            lines = lines[1:]
                        if lines and lines[-1].strip() == "```":
                            lines = lines[:-1]
                    result_text = "\n".join(lines).strip()
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
                return result_text
            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
            return str(result_text)

        except subprocess.TimeoutExpired:
            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 超时({timeout}秒)")
            if attempt < max_retries - 1:
                continue
            raise RuntimeError(f"LLM 调用超时（{timeout}秒），已重试{max_retries}次")
        except FileNotFoundError:
            raise RuntimeError("Claude Code CLI 未找到")

    raise RuntimeError("LLM 调用失败，已达到最大重试次数")


# =============================================================================
# 统一工作流引擎
# =============================================================================

from src.workflow import (
    Coordinator,
    DependencyType,
    MathModelingTemplate,
    CourseworkTemplate,
    FinancialAnalysisTemplate,
    get_template,
    CritiqueEngine,
    CodeExecutor,
    PaperGenerator,
)
from src.document_processing import DocumentLoader


try:
    from src.knowledge import KnowledgeBase
    _KNOWLEDGE_AVAILABLE = True
except ImportError:
    _KNOWLEDGE_AVAILABLE = False


class UnifiedWorkflow:
    """
    统一工作流引擎 v2.0

    四阶段架构（借鉴 LLM-MM-Agent）：
    Stage 1: Problem Analysis - 问题分析 + DAG构建
    Stage 2: Mathematical Modeling - 数学建模（按DAG拓扑序）
    Stage 3: Computational Solving - 计算求解（Claude CLI代码）
    Stage 4: Paper Generation - 论文生成（大纲驱动 + 相关性过滤）
    """

    def __init__(
        self,
        output_dir: str = "work",
        template_name: str = "math_modeling",
        use_knowledge_base: bool = True,
        use_critique: bool = True,
    ):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_name = template_name
        self.use_critique = use_critique
        self.use_knowledge_base = use_knowledge_base and _KNOWLEDGE_AVAILABLE

        # 初始化各组件
        self.coordinator = Coordinator()
        self.critique_engine = CritiqueEngine(_call_llm)
        self.code_executor = CodeExecutor(
            call_llm=_call_llm,
            output_dir=str(self.output_dir),
        )
        self.paper_generator = PaperGenerator(
            call_llm=_call_llm,
            template=get_template(template_name),
            output_dir=str(self.output_dir),
        )
        self.document_loader = DocumentLoader()
        self.knowledge_base = KnowledgeBase() if self.use_knowledge_base else None

        # 全局上下文
        self.context: Dict[str, Any] = {}
        self.problem_text = ""
        self.data_files: Dict[str, str] = {}

        # ======================================================================
        # 显式记忆池：每个阶段完成后生成结构化摘要，供下一阶段调用
        # ======================================================================
        self.memory_pool: Dict[str, str] = {
            "analysis_summary": "",      # 阶段1摘要
            "modeling_summary": "",      # 阶段2摘要
            "algorithm_summary": "",     # 阶段3算法摘要
            "results_summary": "",       # 阶段3结果摘要
            "chapter_summaries": {},     # 论文各章摘要，key=chapter_id
        }

    def run_full_workflow(
        self,
        problem_text: str,
        data_files: Dict[str, str],
        problem_name: str = "数学建模问题",
    ) -> str:
        """
        运行完整工作流（分段生成 + 显式记忆传递）

        核心改进：
        1. 每个阶段完成后，调用LLM生成结构化摘要，存入 memory_pool
        2. 下一阶段只接收摘要，而非原始长文本，避免上下文过载
        3. 论文生成阶段，先预生成大纲，再逐章生成，每章基于前几章摘要 + 阶段摘要
        """
        print("\n" + "=" * 70)
        print("数学建模论文自动生成系统 v2.1")
        print("架构: 分段生成 + 显式记忆池")
        print(f"模板: {self.template_name}")
        print("=" * 70)

        self.problem_text = problem_text
        self.data_files = data_files
        self.context["problem_text"] = problem_text
        self.context["data_files"] = data_files

        # Stage 1: 问题分析
        print(f"\n{'='*60}")
        print("Stage 1: 问题分析")
        print(f"{'='*60}")
        analysis = self._stage_problem_analysis()
        self.context["analysis"] = analysis
        # 生成阶段1摘要
        self.memory_pool["analysis_summary"] = self._summarize_analysis(analysis)
        self._save_text("stage_1_analysis/analysis_summary.md", self.memory_pool["analysis_summary"])

        # Stage 2: 数学建模
        print(f"\n{'='*60}")
        print("Stage 2: 数学建模")
        print(f"{'='*60}")
        modeling = self._stage_mathematical_modeling(analysis)
        self.context["modeling"] = modeling
        # 生成阶段2摘要
        self.memory_pool["modeling_summary"] = self._summarize_modeling(modeling)
        self._save_text("stage_2_modeling/modeling_summary.md", self.memory_pool["modeling_summary"])

        # Stage 3: 计算求解
        print(f"\n{'='*60}")
        print("Stage 3: 计算求解")
        print(f"{'='*60}")
        solving = self._stage_computational_solving(modeling)
        self.context["execution_result"] = solving.get("execution_result", {})
        self.context["code"] = solving.get("code", "")
        self.context["result_analysis"] = solving.get("interpretation", "")
        # 生成阶段3摘要
        self.memory_pool["algorithm_summary"] = self._summarize_algorithm(solving)
        self.memory_pool["results_summary"] = self._summarize_results(solving)
        self._save_text("stage_3_algorithm/algorithm_summary.md", self.memory_pool["algorithm_summary"])
        self._save_text("stage_6_result_analysis/results_summary.md", self.memory_pool["results_summary"])

        # Stage 4: 论文生成（分段 + 记忆传递）
        print(f"\n{'='*60}")
        print("Stage 4: 论文生成（分段逐章 + 记忆衔接）")
        print(f"{'='*60}")
        paper = self._stage_paper_generation_v2()

        # 保存结果
        paper_path = self.output_dir / "final" / "MathModeling_Paper.md"
        paper_path.parent.mkdir(parents=True, exist_ok=True)
        paper_path.write_text(paper, encoding="utf-8")

        # 字数统计
        chinese_chars = len(re.findall(r'[一-鿿]', paper))
        print(f"\n{'='*60}")
        print("论文生成完成")
        print(f"{'='*60}")
        print(f"论文文件: {paper_path}")
        print(f"总字符数: {len(paper)}")
        print(f"中文字数: {chinese_chars}")

        # 导出解决方案
        solution_path = self.output_dir / "final" / "solution.json"
        self.coordinator.export_solution(solution_path)

        # 导出 docx
        self._convert_to_docx(paper_path)

        # 导出记忆池
        self._save_json("final/memory_pool.json", self.memory_pool)

        return paper

    # ========================================================================
    # Stage 1: 问题分析
    # ========================================================================

    def _stage_problem_analysis(self) -> Dict[str, Any]:
        """
        问题分析阶段

        借鉴 LLM-MM-Agent 的反思链设计 + Critique-Improvement
        """
        print("  分析问题并构建DAG...")

        # 加载数据文件摘要
        data_descriptions = self._load_data_descriptions()
        self.context["data_descriptions"] = data_descriptions

        # 注册知识库（如果有相关文档）
        if self.knowledge_base and data_descriptions:
            for name, desc in data_descriptions.items():
                self.knowledge_base.add_document(
                    title=f"数据文件: {name}",
                    content=desc,
                    metadata={"type": "data"},
                )

        # Prompt: 反思链设计（借鉴 LLM-MM-Agent PROBLEM_ANALYSIS_PROMPT）
        prompt = f"""请对以下数学建模问题进行深度分析。

【赛题】
{self.problem_text[:4000]}

【数据文件描述】
{chr(10).join(data_descriptions.values())[:2000]}

要求：
1. 识别问题的核心组件和它们之间的依赖关系
2. 分析问题的动态特性（时间演化、空间分布等）
3. 从多个视角审视问题（数学视角、物理视角、工程视角）
4. 识别关键假设和潜在的不确定性
5. 将问题分解为可管理的子任务

输出严格的JSON格式：
{{
  "background": "问题背景概述（300字）",
  "sub_problems": [
    {{
      "id": "task_1",
      "description": "子任务描述",
      "objective": "任务目标",
      "key_constraints": ["约束1", "约束2"],
      "suggested_methods": ["建议方法1", "建议方法2"]
    }}
  ],
  "key_assumptions": ["假设1", "假设2"],
  "data_dependencies": {{"task_1": ["数据文件1"]}},
  "solution_approach": "整体解决思路（500字）"
}}"""

        # 生成 + Critique-Improvement
        print("  生成初始分析...")
        result = _call_llm(prompt, "你是一位数学建模专家，擅长深度问题分析。")

        analysis = self._parse_json_safely(result, self._default_analysis())

        if self.use_critique:
            print("  启动Critique-Improvement循环...")
            analysis_text = json.dumps(analysis, ensure_ascii=False, indent=2)
            improved = self.critique_engine.critique_and_improve(
                content=analysis_text,
                content_type="analysis",
                context=self.problem_text[:2000],
                max_iterations=1,
                score_threshold=7.5,
            )
            analysis = self._parse_json_safely(improved, analysis)

        # 构建DAG
        sub_problems = analysis.get("sub_problems", [])
        for i, sp in enumerate(sub_problems):
            task_id = sp.get("id", f"task_{i+1}")
            deps = {}
            # 简单的顺序依赖：后面的任务依赖前面的
            if i > 0:
                prev_id = sub_problems[i-1].get("id", f"task_{i}")
                deps[prev_id] = [DependencyType.STRUCTURAL, DependencyType.DATA]
            self.coordinator.register_task(
                task_id=task_id,
                description=sp.get("description", ""),
                dependencies=deps,
            )

        self.coordinator.analyze_dependencies()
        self.context["sub_problems"] = sub_problems

        # 保存
        self._save_json("stage_1_analysis/analysis.json", analysis)
        return analysis

    def _load_data_descriptions(self) -> Dict[str, str]:
        """加载数据文件描述"""
        descriptions = {}
        for name, filepath in self.data_files.items():
            path = Path(filepath)
            if not path.exists():
                descriptions[name] = f"{name}: 文件不存在"
                continue

            try:
                result = self.document_loader.load(filepath)
                if result.success and result.document:
                    desc = result.document.text[:1500]
                    descriptions[name] = f"【{name}】\n{desc}\n"
                    # 添加到知识库
                    if self.knowledge_base:
                        self.knowledge_base.add_document(
                            title=f"数据: {name}",
                            content=desc,
                            source=filepath,
                            metadata={"type": "data_file", "filename": name},
                        )
                else:
                    descriptions[name] = f"{name}: 加载失败 - {result.error}"
            except Exception as e:
                descriptions[name] = f"{name}: 错误 - {str(e)}"

        return descriptions

    # ========================================================================
    # Stage 2: 数学建模
    # ========================================================================

    def _stage_mathematical_modeling(self, analysis: Dict) -> Dict[str, Any]:
        """
        数学建模阶段 - 按DAG拓扑序执行每个子任务

        借鉴 LLM-MM-Agent：
        - 依赖上下文自动拼接
        - 方法检索增强生成
        - Actor-Critic-Improvement 循环
        """
        sub_problems = analysis.get("sub_problems", [])
        if not sub_problems:
            sub_problems = [{"id": "task_1", "description": "建立模型求解"}]

        # 处理所有子任务的详细建模（确保不遗留pending任务）
        tasks_to_model = self.coordinator.dag_order
        print(f"  共 {len(tasks_to_model)} 个子任务，将全部进行详细建模")

        all_formulas = []
        all_models = []

        for task_id in tasks_to_model:
            task_node = self.coordinator.tasks.get(task_id)
            if not task_node:
                continue

            print(f"\n  建模任务: {task_id} - {task_node.description[:60]}...")

            # 获取依赖上下文
            dep_context = self.coordinator.get_dependency_context(
                task_id, max_chars=3000
            )

            # 知识库检索相关方法
            kb_context = ""
            if self.knowledge_base:
                kb_results = self.knowledge_base.query_with_context(
                    task_node.description, top_k=3, max_chars=1500
                )
                if kb_results:
                    kb_context = f"【相关知识】\n{kb_results}\n\n"

            # 1. 任务分析（缩短prompt以加快响应）
            analysis_prompt = f"""对以下子任务进行数学分析：

任务: {task_node.description}

背景: {self.problem_text[:1500]}

{dep_context}

{kb_context}

分析核心数学结构、建模方法、关键变量。输出至少600字。"""

            task_analysis = _call_llm(
                analysis_prompt,
                "你是数学建模专家，擅长将实际问题抽象为数学结构。"
            )

            # 2. 公式生成（Actor）
            formulas_prompt = f"""基于以下分析建立数学模型。

任务分析:
{task_analysis[:1200]}

要求：
1. 定义变量和参数（表格形式）
2. 建立核心数学公式（LaTeX格式，公式编号）
3. 说明公式物理意义
4. 列出模型假设

输出至少1200字的完整建模过程。"""

            formulas = _call_llm(
                formulas_prompt,
                "你是数学建模专家，擅长严谨的形式化建模。"
            )

            # 3. Critique-Improvement
            if self.use_critique:
                print(f"    对公式进行Critique-Improvement...")
                formulas = self.critique_engine.critique_and_improve(
                    content=formulas,
                    content_type="modeling",
                    context=f"任务: {task_node.description}\n{self.problem_text[:1000]}",
                    max_iterations=1,
                    score_threshold=7.5,
                    min_chars=1500,
                )

            # 保存结果
            self.coordinator.save_task_result(
                task_id, {"analysis": task_analysis, "formulas": formulas}, key="modeling"
            )
            all_formulas.append(formulas)
            all_models.append(task_analysis)

        modeling_result = {
            "formulas": "\n\n".join(all_formulas),
            "models": "\n\n".join(all_models),
        }

        self.context["formulas"] = modeling_result["formulas"]
        self._save_json("stage_2_modeling/modeling.json", modeling_result)
        self._save_text("stage_2_modeling/formulas.md", modeling_result["formulas"])

        return modeling_result

    # ========================================================================
    # Stage 3: 计算求解
    # ========================================================================

    def _stage_computational_solving(self, modeling: Dict) -> Dict[str, Any]:
        """
        计算求解阶段

        明确使用 Claude CLI 生成代码
        使用 CodeExecutor 执行与调试
        """
        print("  [Stage 3] 开始计算求解...")

        formulas = modeling.get("formulas", "")

        # 设计算法
        print("  [Stage 3.1] 设计求解算法...")
        algorithm_desc = self._design_algorithm(modeling)
        self.context["algorithm"] = algorithm_desc
        print(f"  [Stage 3.1] 算法设计完成 ({len(algorithm_desc)} 字符)")

        # 构建代码生成 Prompt
        # 注：要求代码精简，避免输出被截断
        code_prompt = f"""请编写精简的Python求解代码，解决以下数学建模问题。

【数学模型核心】
{formulas[:1500]}

【算法设计】
{algorithm_desc[:1000]}

【赛题摘要】
{self.problem_text[:1000]}

【数据文件】
{json.dumps(self.data_files, ensure_ascii=False, indent=2)}

硬性要求：
1. 代码必须精简（不超过200行），不要定义不必要的类，用函数即可
2. 代码开头必须是 import 语句，不要有任何中文说明文字
3. 代码必须能够读取上述数据文件（使用 pandas 读取 Excel）
4. 代码运行后必须使用以下代码写入结果：
   import os, json
   output_dir = os.environ.get('OUTPUT_DIR', '.')
   out_path = os.path.join(output_dir, 'execution', 'results.json')
   os.makedirs(os.path.dirname(out_path), exist_ok=True)
   with open(out_path, 'w') as f:
       json.dump(results, f, indent=2)
5. 代码中必须包含 main() 函数，且 if __name__ == '__main__': main()
6. 所有数值结果使用 Python 原生 float/int，不要嵌套numpy类型
7. 不得使用 input() 等交互式函数
8. 结果必须包含具体的数值，不能是空值或占位符
9. 代码必须在60秒内运行完成。禁止使用大规模网格搜索或复杂全局优化。如需搜索，网格点总数不得超过200个。
10. 优先使用解析解或简化公式，避免复杂数值积分和迭代优化。

输出纯Python代码，不要包含 markdown 代码块标记，不要任何解释。"""

        # 使用 CodeExecutor 生成代码
        # 注：代码生成使用 API（输出更稳定、长度限制宽松），代码修复使用 Claude CLI
        print("  [Stage 3.2] 生成求解代码...")
        result = self.code_executor.generate_and_run(
            prompt=code_prompt,
            system_prompt="你是Python编程专家，只输出代码，不输出任何解释。",
            data_files=self.data_files,
            filename="solve.py",
            use_claude_cli=False,
        )
        print(f"  [Stage 3.2] 代码生成完成，success={result.get('success', False)}")

        code = result.get("code", "")
        execution_result = result.get("execution_result", {})

        # 结果解读
        print("  解读计算结果...")
        interpretation = self._interpret_results(execution_result, modeling)

        # 保存
        self._save_text("stage_4_coding/solve.py", code)
        self._save_json("stage_5_execution/execution_result.json", execution_result)
        self._save_text("stage_6_result_analysis/interpretation.md", interpretation)

        return {
            "code": code,
            "execution_result": execution_result,
            "interpretation": interpretation,
            "success": result.get("success", False),
        }

    def _design_algorithm(self, modeling: Dict) -> str:
        """设计求解算法"""
        formulas = modeling.get("formulas", "")

        prompt = f"""基于以下数学模型，设计详细的求解算法。

【数学模型】
{formulas[:2000]}

要求：
1. 选择合适的数值方法或优化算法
2. 给出详细的算法步骤（伪代码或分步说明）
3. 分析算法的时间复杂度和空间复杂度
4. 讨论算法的收敛性和稳定性
5. 设置关键参数及其选择依据

输出至少800字的算法设计文档。"""

        print("    调用LLM设计算法...")
        return _call_llm(prompt, "你是算法设计专家。")

    def _interpret_results(self, execution_result: Dict, modeling: Dict) -> str:
        """解读计算结果"""
        result_text = json.dumps(execution_result, ensure_ascii=False, indent=2)
        print(f"    计算结果大小: {len(result_text)} 字符")

        prompt = f"""请对以下计算结果进行专业解读。

【计算结果】
{result_text[:3000]}

【数学模型】
{modeling.get('formulas', '')[:1500]}

要求：
1. 总结主要数值结果
2. 分析结果的合理性和物理意义
3. 进行误差分析（如有可能）
4. 讨论结果的稳健性
5. 给出明确的研究结论

输出至少1500字的详细解读。"""

        return _call_llm(prompt, "你是数据分析专家。")

    # ========================================================================
    # Stage 4: 论文生成
    # ========================================================================

    def _stage_paper_generation_v2(self) -> str:
        """
        论文生成阶段 v2.1（分段生成 + 显式记忆传递）

        核心改进：
        1. 预生成论文章节大纲（每章列出必须覆盖的要点）
        2. 逐章生成时，prompt 中明确包含：本章大纲 + 相关阶段摘要 + 前几章摘要
        3. 每章生成后，生成该章结构化摘要，存入记忆池，供后续章节衔接
        """
        print("  [Stage 4] 开始论文生成（分段逐章 + 记忆衔接）...")

        # 生成图表
        print("  [Stage 4.1] 生成图表...")
        charts = self._generate_charts()
        self.context["charts"] = charts

        # 组装记忆池给 PaperGenerator
        memory_for_paper = {
            "analysis_summary": self.memory_pool.get("analysis_summary", ""),
            "modeling_summary": self.memory_pool.get("modeling_summary", ""),
            "algorithm_summary": self.memory_pool.get("algorithm_summary", ""),
            "results_summary": self.memory_pool.get("results_summary", ""),
        }

        # 使用 PaperGenerator v2 生成论文
        print("  [Stage 4.2] 预生成大纲，再逐章生成...")
        paper = self.paper_generator.generate_paper_v2(
            context=self.context,
            memory_pool=memory_for_paper,
            use_critique=self.use_critique,
        )

        # 保存章节摘要
        self.memory_pool["chapter_summaries"] = self.paper_generator.chapter_summaries
        self._save_json("final/chapter_summaries.json", self.paper_generator.chapter_summaries)

        # 保存
        paper_path = self.paper_generator.save_paper(paper, "MathModeling_Paper.md")
        print(f"  [Stage 4.3] 论文已保存: {paper_path}")

        return paper

    def _summarize_analysis(self, analysis: Dict) -> str:
        """生成阶段1的结构化摘要（供阶段2/4使用）"""
        analysis_text = json.dumps(analysis, ensure_ascii=False, indent=2)
        prompt = f"""请对以下问题分析结果进行结构化提炼，生成一份"分析摘要"（严格控制在400-500字）。

【原始分析】
{analysis_text[:2000]}

【摘要要求】
1. 问题背景（1句话）
2. 关键子问题（每题1句话，只列目标）
3. 核心假设（3条）
4. 解决思路（2-3句话）
5. 关键数值约束（列出3个最重要的）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数学建模专家，擅长提炼问题要点。")

    def _summarize_modeling(self, modeling: Dict) -> str:
        """生成阶段2的结构化摘要（供阶段3/4使用）"""
        formulas = modeling.get("formulas", "")[:1500]
        prompt = f"""请对以下数学建模结果进行结构化提炼，生成一份"建模摘要"（严格控制在400-500字）。

【原始建模内容】
{formulas}

【摘要要求】
1. 核心变量（列出5个最重要的变量、符号、单位）
2. 核心公式（1-2个最关键的LaTeX公式）
3. 求解策略（1句话）
4. 与子问题的映射（每子问题1句话）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数学建模专家，擅长提炼数学模型要点。")

    def _summarize_algorithm(self, solving: Dict) -> str:
        """生成阶段3算法摘要"""
        code = solving.get("code", "")[:1000]
        prompt = f"""请对以下求解算法进行结构化提炼（严格控制在200-300字）。

【代码片段】
{code}

【摘要要求】
1. 算法名称
2. 输入数据
3. 核心步骤（3步）
4. 输出格式

输出纯文本。"""
        return _call_llm(prompt, "你是算法专家。")

    def _summarize_results(self, solving: Dict) -> str:
        """生成阶段3结果摘要（供论文阶段使用）"""
        result = solving.get("execution_result", {})
        result_text = json.dumps(result, ensure_ascii=False, indent=2)[:1500]
        interpretation = solving.get("interpretation", "")[:1000]

        prompt = f"""请对以下计算结果进行结构化提炼，生成一份"结果摘要"（严格控制在400-500字）。

【计算结果】
{result_text}

【结果解读】
{interpretation}

【摘要要求】
1. 关键数值结果（用表格列出最重要的5个数值）
2. 主要发现（2-3句话）
3. 与模型的对应（1句话）
4. 可视化建议（1句话）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数据分析专家，擅长提炼数值结果。")

    def _generate_charts(self) -> str:
        """生成论文图表"""
        charts_dir = self.output_dir / "stage_7_charts"
        charts_dir.mkdir(parents=True, exist_ok=True)

        execution_result = self.context.get("execution_result", {})

        # 使用 matplotlib 生成基础图表
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import numpy as np
            plt.rcParams['font.sans-serif'] = [
                'Noto Sans CJK SC', 'SimHei', 'DejaVu Sans'
            ]
            plt.rcParams['axes.unicode_minus'] = False

            chart_count = 0

            # 深度提取数值数据
            def extract_numeric_data(data, prefix=""):
                """递归提取字典中的数值数据"""
                keys, vals = [], []
                if isinstance(data, dict):
                    for k, v in data.items():
                        key_path = f"{prefix}_{k}" if prefix else k
                        if isinstance(v, (int, float)) and not isinstance(v, bool):
                            keys.append(key_path)
                            vals.append(float(v))
                        elif isinstance(v, dict):
                            sub_k, sub_v = extract_numeric_data(v, key_path)
                            keys.extend(sub_k)
                            vals.extend(sub_v)
                return keys, vals

            numeric_keys, numeric_vals = extract_numeric_data(execution_result)

            # 策略1：如果有多个数值结果，生成对比柱状图
            if len(numeric_keys) > 1 and len(numeric_vals) > 1:
                # 限制数量避免图表过密
                if len(numeric_keys) > 20:
                    numeric_keys = numeric_keys[:20]
                    numeric_vals = numeric_vals[:20]

                fig, ax = plt.subplots(figsize=(max(10, len(numeric_keys)*0.5), 6))
                colors = plt.cm.viridis(np.linspace(0, 0.8, len(numeric_keys)))
                ax.bar(range(len(numeric_keys)), numeric_vals, color=colors)
                ax.set_xticks(range(len(numeric_keys)))
                ax.set_xticklabels(numeric_keys, rotation=45, ha='right', fontsize=8)
                ax.set_xlabel('指标', fontsize=12)
                ax.set_ylabel('数值', fontsize=12)
                ax.set_title('计算结果数值对比', fontsize=14)
                plt.tight_layout()
                fig.savefig(str(charts_dir / 'fig_01_comparison.png'), dpi=300)
                plt.close(fig)
                chart_count += 1

            # 策略2：如果结果是导弹/无人机相关数据，生成轨迹或参数图
            # 尝试从 results.json 或 execution_result 中提取特殊结构
            results_json_path = self.output_dir / "execution" / "results.json"
            if results_json_path.exists():
                try:
                    with open(results_json_path, "r", encoding="utf-8") as f:
                        results_data = json.load(f)

                    # 如果结果是按问题组织的（如 task_1, task_2），生成饼图
                    if isinstance(results_data, dict):
                        task_results = {}
                        for k, v in results_data.items():
                            if isinstance(v, dict) and "tau_eff" in v:
                                task_results[k] = v["tau_eff"]
                            elif isinstance(v, dict) and "result" in v:
                                try:
                                    task_results[k] = float(v["result"])
                                except:
                                    pass

                        if len(task_results) > 1:
                            fig, ax = plt.subplots(figsize=(8, 8))
                            labels = list(task_results.keys())
                            sizes = list(task_results.values())
                            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
                            ax.set_title('各任务有效遮蔽时间占比', fontsize=14)
                            plt.tight_layout()
                            fig.savefig(str(charts_dir / 'fig_02_pie.png'), dpi=300)
                            plt.close(fig)
                            chart_count += 1
                except Exception as e:
                    print(f"    读取 results.json 生成图表失败: {e}")

            # 策略3：生成流程图/示意图（如果没有任何数值图表）
            if chart_count == 0:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "图表生成依赖于数值计算结果\n\n"
                       "请确保代码执行产生了有效的数值结果",
                       ha='center', va='center', fontsize=14, color='gray')
                ax.set_xlim(0, 1)
                ax.set_ylim(0, 1)
                ax.axis('off')
                fig.savefig(str(charts_dir / 'fig_placeholder.png'), dpi=150)
                plt.close(fig)
                chart_count += 1

            print(f"  生成 {chart_count} 个图表")
        except ImportError:
            print("  matplotlib 未安装，跳过图表生成")

        return f"图表保存于: {charts_dir}"


    # ========================================================================
    # 工具方法
    # ========================================================================

    def _convert_to_docx(self, md_path: Path):
        """将Markdown论文转换为docx格式"""
        docx_path = md_path.parent / "数学建模论文.docx"

        # 方法1: 使用 pandoc（质量最好）
        try:
            import subprocess
            result = subprocess.run(
                ["pandoc", str(md_path), "-o", str(docx_path), "--resource-path", str(md_path.parent)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and docx_path.exists():
                print(f"  论文已导出为Word: {docx_path}")
                return
        except FileNotFoundError:
            pass
        except Exception as e:
            print(f"  pandoc导出失败: {e}")

        # 方法2: 使用 python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            # 设置中文字体支持
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 读取markdown内容
            text = md_path.read_text(encoding="utf-8")

            # 简单的markdown到docx转换
            lines = text.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]

                # 标题
                if line.startswith("# "):
                    p = doc.add_heading(line[2:], level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith("#### "):
                    doc.add_heading(line[5:], level=3)
                # 表格（简单处理）
                elif line.startswith("|"):
                    # 收集表格行
                    table_lines = []
                    while i < len(lines) and lines[i].startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    i -= 1  # 回退一行，外层循环会+1

                    if len(table_lines) >= 2:
                        # 解析表头
                        header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                        # 创建表格
                        table = doc.add_table(rows=1, cols=len(header_cells))
                        table.style = 'Light Grid Accent 1'
                        hdr_cells = table.rows[0].cells
                        for j, cell_text in enumerate(header_cells):
                            hdr_cells[j].text = cell_text
                        # 添加数据行（跳过第二行的分隔符）
                        for row_line in table_lines[2:]:
                            cells = [c.strip() for c in row_line.split("|")[1:-1]]
                            if len(cells) == len(header_cells):
                                row_cells = table.add_row().cells
                                for j, cell_text in enumerate(cells):
                                    row_cells[j].text = cell_text
                # 图片
                elif line.startswith("!["):
                    # 尝试提取图片路径
                    match = re.search(r'!\[.*?\]\((.*?)\)', line)
                    if match:
                        img_path = md_path.parent / match.group(1)
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.5))
                            except Exception:
                                pass
                # 普通段落
                elif line.strip():
                    doc.add_paragraph(line)

                i += 1

            doc.save(str(docx_path))
            print(f"  论文已导出为Word: {docx_path}")
        except ImportError:
            print("  未安装 python-docx，跳过Word导出。可运行: pip install python-docx")
        except Exception as e:
            print(f"  Word导出失败: {e}")

    def _save_json(self, rel_path: str, data: Any):
        """保存JSON文件"""
        path = self.output_dir / rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _save_text(self, rel_path: str, text: str):
        """保存文本文件"""
        path = self.output_dir / rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding="utf-8")

    def _parse_json_safely(self, text: str, default: Any) -> Any:
        """安全解析JSON"""
        try:
            # 尝试提取 JSON 块
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                return json.loads(match.group())
            return json.loads(text)
        except Exception:
            return default

    def _default_analysis(self) -> Dict:
        """默认分析问题结果"""
        return {
            "background": "数学建模问题分析",
            "sub_problems": [
                {
                    "id": "task_1",
                    "description": "建立模型求解",
                    "objective": "解答问题",
                    "key_constraints": [],
                    "suggested_methods": ["数学建模方法"],
                }
            ],
            "key_assumptions": ["合理假设"],
            "data_dependencies": {},
            "solution_approach": "建立数学模型进行求解",
        }


def run_auto_paper_generation(
    problem_file: str = "problem.md",
    data_files: Dict[str, str] = None,
    output_dir: str = "work",
    template_name: str = "math_modeling",
) -> str:
    """运行全自动论文生成"""
    if data_files is None:
        data_files = {}

    problem_text = ""
    if Path(problem_file).exists():
        with open(problem_file, "r", encoding="utf-8") as f:
            problem_text = f.read()

    engine = UnifiedWorkflow(
        output_dir=output_dir,
        template_name=template_name,
    )

    return engine.run_full_workflow(
        problem_text=problem_text,
        data_files=data_files,
    )
